"""Load supported user documents for indexing."""

from __future__ import annotations

import io
import logging
from collections.abc import Iterable
from dataclasses import dataclass
from pathlib import Path
from typing import BinaryIO

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


@dataclass(slots=True)
class LoadedDocument:
    doc_id: str
    filename: str
    text: str


def _normalize_line_endings(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _read_text_from_bytes(raw: bytes, filename: str = "") -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            text = _normalize_line_endings(raw.decode(encoding))
            logger.debug("Decoded %s with encoding=%s", filename, encoding)
            return text
        except UnicodeDecodeError:
            continue
    logger.warning("All encodings failed for %s; using UTF-8 with replacement", filename)
    return _normalize_line_endings(raw.decode("utf-8", errors="replace"))


def _load_pdf(raw: bytes, filename: str) -> str:
    import fitz  # pymupdf

    doc = fitz.open(stream=raw, filetype="pdf")
    pages = [page.get_text() for page in doc]
    doc.close()
    text = "\n\n".join(pages)
    logger.info("PDF loaded: %s (%d pages, %d chars)", filename, len(pages), len(text))
    return _normalize_line_endings(text)


def _load_docx(raw: bytes, filename: str) -> str:
    import docx

    doc = docx.Document(io.BytesIO(raw))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    text = "\n\n".join(paragraphs)
    logger.info("DOCX loaded: %s (%d paragraphs, %d chars)", filename, len(paragraphs), len(text))
    return _normalize_line_endings(text)


def load_documents(files: Iterable[BinaryIO]) -> list[LoadedDocument]:
    """Load uploaded documents into memory.

    Supports .txt, .md, .pdf and .docx.
    Expects file-like objects with a ``name`` attribute and ``read()`` method.
    """
    docs: list[LoadedDocument] = []
    for file in files:
        raw_name = getattr(file, "name", "unknown")
        filename = Path(raw_name).name
        suffix = Path(filename).suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            logger.warning("Skipping unsupported file: %s (suffix=%s)", filename, suffix)
            continue

        logger.debug("Loading file: %s (raw path: %s)", filename, raw_name)
        raw = file.read()
        if isinstance(raw, str):
            raw = raw.encode("utf-8")

        if suffix == ".pdf":
            text = _load_pdf(raw, filename)
        elif suffix == ".docx":
            text = _load_docx(raw, filename)
        else:
            text = _read_text_from_bytes(raw, filename)

        logger.info("Loaded document: %s (%d chars)", filename, len(text))
        docs.append(LoadedDocument(doc_id=filename, filename=filename, text=text))
    logger.info("load_documents: %d document(s) loaded", len(docs))
    return docs
